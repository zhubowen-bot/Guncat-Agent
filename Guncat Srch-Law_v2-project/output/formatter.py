"""
Output Formatter - 法律意见书格式化输出模块
支持 Markdown / Word / PDF 三种输出格式
"""

from typing import Dict, Any, List, Optional
import os
from pathlib import Path
from datetime import datetime


class LegalOpinionFormatter:
    """
    法律意见书格式化器

    功能：
    1. 将分析结果格式化为结构化法律意见书
    2. 支持 Markdown / Word / PDF 输出
    3. 严格遵循输出模板
    """

    def __init__(self, template_path: Optional[str] = None):
        """
        初始化格式化器

        Args:
            template_path: 模板文件路径（Markdown格式）
        """
        self.template_path = template_path
        self.template = self._load_template()

    def _load_template(self) -> str:
        """加载输出模板"""
        if self.template_path and os.path.exists(self.template_path):
            with open(self.template_path, "r", encoding="utf-8") as f:
                return f.read()
        else:
            # 使用内置默认模板
            return self._default_template()

    def _default_template(self) -> str:
        """默认输出模板"""
        return """# 法律意见书

## 📋 摘要

- **案件概述**：{summary}
- **涉及主体**：{parties}
- **核心争议点**：{key_issues}
- **分析日期**：{date}

## 🔑 关键术语辨析

{term_analysis}

## ⚖️ 法律适用分析

### 现行有效法律依据

{laws_table}

### 司法解释与监管文件

{regulations_table}

## 🏛️ 国企合规专项分析

{compliance_analysis}

## 📜 合同体系化解释

{contract_analysis}

## 🧠 法理深度分析

{legal_analysis}

## ⚖️ 刑事风险评估

{criminal_analysis}

## ⚠️ 风险识别与等级评估

{risk_table}

## 🎯 法律意见与建议

### 即时应对措施

{immediate_actions}

### 证据保全与财产保全建议

{preservation_advice}

### 中长期合规建议

{long_term_advice}

### 争议解决路径

{dispute_resolution}

## 📚 参考案例与裁判思路

{reference_cases}

## ⚡ 时效性提示

{timeliness_note}

---
**保密提醒**：本法律意见仅供内部决策参考，涉及敏感信息请注意保密。
**分析日期**：{date}
**系统版本**：Guncat Srch-Law V2
"""

    def format(
        self,
        analysis_result: Dict[str, Any],
        output_format: str = "markdown",
    ) -> str:
        """
        格式化分析结果

        Args:
            analysis_result: 分析结果字典（由各Agent输出汇总）
            output_format: 输出格式（"markdown" / "docx" / "pdf"）

        Returns:
            格式化后的法律意见书内容（字符串）
        """
        # 提取各模块内容
        context = self._extract_context(analysis_result)

        # 渲染模板
        output = self.template.format(**context)

        return output

    def _extract_context(self, result: Dict[str, Any]) -> Dict[str, str]:
        """从分析结果中提取上下文变量"""
        # 默认值
        ctx = {
            "summary": "（待填写）",
            "parties": "（待填写）",
            "key_issues": "（待填写）",
            "date": datetime.now().strftime("%Y年%m月%d日"),
            "term_analysis": "（待补充）",
            "laws_table": "（待补充）",
            "regulations_table": "（待补充）",
            "compliance_analysis": "（待补充）",
            "contract_analysis": "（待补充）",
            "legal_analysis": "（待补充）",
            "criminal_analysis": "（待补充）",
            "risk_table": "（待补充）",
            "immediate_actions": "（待补充）",
            "preservation_advice": "（待补充）",
            "long_term_advice": "（待补充）",
            "dispute_resolution": "（待补充）",
            "reference_cases": "（待补充）",
            "timeliness_note": "（无特殊时效性提示）",
        }

        # 从结果中提取（根据实际输出结构调整）
        if isinstance(result, dict):
            # 如果有摘要
            if "summary" in result:
                ctx["summary"] = result["summary"]
            # 如果有完整文本输出，尝试智能分段
            if "result" in result and isinstance(result["result"], str):
                full_text = result["result"]
                # 尝试按标题分段（简单启发式）
                ctx = self._auto_section(full_text, ctx)

        return ctx

    def _auto_section(self, text: str, ctx: Dict[str, str]) -> Dict[str, str]:
        """尝试自动识别文本中的章节，填入对应字段"""
        # 简单的关键词匹配分段
        if "术语辨析" in text or "关键术语" in text:
            ctx["term_analysis"] = self._extract_section(text, ["术语辨析", "关键术语"])
        if "法律适用" in text or "法律依据" in text:
            ctx["laws_table"] = self._extract_section(text, ["法律适用", "法律依据"])
        if "国企" in text or "合规" in text:
            ctx["compliance_analysis"] = self._extract_section(text, ["国企", "合规"])
        if "合同" in text:
            ctx["contract_analysis"] = self._extract_section(text, ["合同"])
        if "法理" in text:
            ctx["legal_analysis"] = self._extract_section(text, ["法理"])
        if "刑事" in text or "犯罪" in text:
            ctx["criminal_analysis"] = self._extract_section(text, ["刑事", "犯罪"])
        if "风险" in text:
            ctx["risk_table"] = self._extract_section(text, ["风险"])
        if "建议" in text or "措施" in text:
            ctx["immediate_actions"] = self._extract_section(text, ["建议", "措施"])
        if "案例" in text:
            ctx["reference_cases"] = self._extract_section(text, ["案例"])

        return ctx

    def _extract_section(self, text: str, keywords: List[str]) -> str:
        """根据关键词从文本中提取相关段落（简化版）"""
        lines = text.split("\n")
        relevant = []
        for line in lines:
            for kw in keywords:
                if kw in line:
                    relevant.append(line)
                    break
        return "\n".join(relevant[:10]) if relevant else "（待补充）"

    def save(
        self,
        content: str,
        output_path: str,
        output_format: str = "markdown",
    ) -> str:
        """
        保存法律意见书到文件

        Args:
            content: 格式化后的内容
            output_path: 输出文件路径（不含扩展名）
            output_format: 输出格式

        Returns:
            实际保存的文件路径
        """
        if output_format == "markdown":
            return self._save_markdown(content, output_path)
        elif output_format == "docx":
            return self._save_docx(content, output_path)
        elif output_format == "pdf":
            return self._save_pdf(content, output_path)
        else:
            raise ValueError(f"不支持的输出格式: {output_format}")

    def _save_markdown(self, content: str, output_path: str) -> str:
        """保存为Markdown文件"""
        full_path = f"{output_path}.md"
        with open(full_path, "w", encoding="utf-8") as f:
            f.write(content)
        return full_path

    def _save_docx(self, content: str, output_path: str) -> str:
        """保存为Word文档（.docx）"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 设置文档样式
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)

            # 将Markdown内容转换为Word（简化版，实际使用可接入专门库）
            lines = content.split("\n")
            for line in lines:
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    doc.add_paragraph(line)

            full_path = f"{output_path}.docx"
            doc.save(full_path)
            return full_path

        except ImportError:
            print("警告：未安装 python-docx，无法生成 Word 文档。请运行 pip install python-docx")
            # 降级为保存Markdown
            return self._save_markdown(content, output_path)

    def _save_pdf(self, content: str, output_path: str) -> str:
        """保存为PDF文件"""
        try:
            import subprocess

            # 先保存为Markdown
            md_path = self._save_markdown(content, output_path)

            # 尝试使用 weasyprint 或 pandoc 转换为PDF
            try:
                from weasyprint import HTML
                pdf_path = f"{output_path}.pdf"
                HTML(filename=md_path).write_pdf(pdf_path)
                return pdf_path
            except ImportError:
                # 尝试使用 pandoc
                try:
                    subprocess.run(
                        ["pandoc", md_path, "-o", f"{output_path}.pdf"],
                        check=True,
                        capture_output=True,
                    )
                    return f"{output_path}.pdf"
                except (subprocess.CalledProcessError, FileNotFoundError):
                    print("警告：未安装 weasyprint 或 pandoc，无法生成 PDF。请安装其中之一。")
                    return md_path

        except Exception as e:
            print(f"生成 PDF 失败: {e}")
            return self._save_markdown(content, output_path)


if __name__ == "__main__":
    # 测试
    formatter = LegalOpinionFormatter()
    test_result = {
        "result": "# 测试分析报告\n## 摘要\n这是一个测试案件。\n## 风险分析\n风险等级：高"
    }
    output = formatter.format(test_result)
    print(output[:500])
